from pathlib import Path

import fitz
import pytest
from docx import Document

from app.services.parser import DocumentParser


@pytest.mark.asyncio
async def test_parse_native_pdf(tmp_path: Path):
    pdf_path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Page one text")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Page two text")
    doc.save(pdf_path)
    doc.close()

    parser = DocumentParser()
    result = parser.parse(pdf_path)
    assert result.page_count == 2
    assert "Page one text" in result.raw_text


@pytest.mark.asyncio
async def test_parse_scanned_pdf(tmp_path: Path, monkeypatch):
    pdf_path = tmp_path / "scanned.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.save(pdf_path)
    doc.close()

    parser = DocumentParser()
    monkeypatch.setattr(parser.ocr_service, "extract_text_per_page", lambda pages: ["OCR PAGE TEXT"])
    result = parser.parse(pdf_path)
    assert "OCR PAGE TEXT" in result.raw_text


@pytest.mark.asyncio
async def test_parse_docx(tmp_path: Path):
    path = tmp_path / "sample.docx"
    d = Document()
    d.add_heading("Master Service Agreement", level=1)
    d.add_paragraph("Payment is due in 30 days.")
    d.save(path)

    parser = DocumentParser()
    result = parser.parse(path)
    assert "Master Service Agreement" in result.raw_text


@pytest.mark.asyncio
async def test_chunk_size(tmp_path: Path):
    path = tmp_path / "chunk.docx"
    d = Document()
    d.add_paragraph("A" * 5000)
    d.save(path)

    parser = DocumentParser()
    result = parser.parse(path)
    assert all(len(chunk) <= 1500 for chunk in result.chunks)
