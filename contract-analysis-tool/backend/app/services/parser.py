from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import fitz
import xlrd
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openpyxl import load_workbook
from PIL import Image

from app.config import get_settings
from app.services.ocr import OCRService


@dataclass
class ParseResult:
    raw_text: str
    chunks: list[str]
    page_count: int
    is_scanned: bool
    metadata: dict


class DocumentParser:
    """
    Converts PDF, DOCX, and Excel files into clean text chunks suitable for LLM
    processing.
    """

    def __init__(self):
        settings = get_settings()
        self.ocr_service = OCRService(settings.tesseract_cmd)
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=2200,
            chunk_overlap=120,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def parse(self, file_path: str | Path, filename: str | None = None, mime_type: str | None = None) -> ParseResult:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            raw_text, page_count, is_scanned = self._parse_pdf(path)
        elif suffix == ".docx":
            raw_text, page_count, is_scanned = self._parse_docx(path), 1, False
        elif suffix == ".xlsx":
            raw_text, page_count, is_scanned = self._parse_xlsx(path), 1, False
        elif suffix == ".xls":
            raw_text, page_count, is_scanned = self._parse_xls(path), 1, False
        else:
            raise ValueError("Unsupported file type. Use PDF, DOCX, XLSX, or XLS.")

        chunks = self.splitter.split_text(raw_text)
        metadata = {
            "filename": filename or path.name,
            "size": path.stat().st_size,
            "mime_type": mime_type or self._guess_mime_type(suffix),
            "page_count": page_count,
        }

        return ParseResult(
            raw_text=raw_text,
            chunks=chunks,
            page_count=page_count,
            is_scanned=is_scanned,
            metadata=metadata,
        )

    def _parse_pdf(self, path: Path) -> tuple[str, int, bool]:
        doc = fitz.open(path)
        texts: list[str] = []
        scanned_pages: list[Image.Image] = []

        for page in doc:
            text = page.get_text("text").strip()
            if text:
                texts.append(text)
                continue
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            scanned_pages.append(img)
            texts.append("")

        is_scanned = len(scanned_pages) > 0 and all(not t for t in texts)
        if scanned_pages:
            ocr_texts = self.ocr_service.extract_text_per_page(scanned_pages)
            ocr_iter = iter(ocr_texts)
            for idx, value in enumerate(texts):
                if not value:
                    texts[idx] = next(ocr_iter, "")

        raw_text = "\n\n".join(texts)
        page_count = len(doc)
        doc.close()
        return raw_text, page_count, is_scanned

    def _parse_docx(self, path: Path) -> str:
        doc = Document(path)
        lines: list[str] = []

        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                if p.style and p.style.name and "Heading" in p.style.name:
                    lines.append(f"\n## {txt}\n")
                else:
                    lines.append(txt)

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    lines.append(" | ".join(row_cells))

        return "\n".join(lines)

    def _parse_xlsx(self, path: Path) -> str:
        wb = load_workbook(filename=path, read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in wb.worksheets:
            lines.append(f"\n## Sheet: {sheet.title}\n")
            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if values:
                    lines.append(" | ".join(values))
        wb.close()
        return "\n".join(lines)

    def _parse_xls(self, path: Path) -> str:
        wb = xlrd.open_workbook(path)
        lines: list[str] = []
        for sheet in wb.sheets():
            lines.append(f"\n## Sheet: {sheet.name}\n")
            for row_idx in range(sheet.nrows):
                row_values: list[str] = []
                for col_idx in range(sheet.ncols):
                    cell = sheet.cell_value(row_idx, col_idx)
                    text = str(cell).strip()
                    if text:
                        row_values.append(text)
                if row_values:
                    lines.append(" | ".join(row_values))
        return "\n".join(lines)

    @staticmethod
    def _guess_mime_type(suffix: str) -> str:
        if suffix == ".pdf":
            return "application/pdf"
        if suffix == ".docx":
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if suffix == ".xlsx":
            return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        if suffix == ".xls":
            return "application/vnd.ms-excel"
        return "application/octet-stream"
